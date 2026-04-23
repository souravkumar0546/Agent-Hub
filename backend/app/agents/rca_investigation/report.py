"""Generate the final investigation report as a .docx buffer.

Matches FORM-GMP-QA-0504 structure. Ported from Devio generation/report.py.
"""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.agents.rca_investigation.session_state import Session


def _add_header(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.cell(0, 0).text = "Syngene"
    table.cell(0, 1).text = "Template"
    table.cell(0, 2).text = "Department:\nQuality Assurance"

    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(1, 0).text = "Reference SOP No.: SOP-GMP-QA-0066"

    table.cell(2, 0).merge(table.cell(2, 2))
    table.cell(2, 0).text = "Title: ANNEXURE 02 - TEMPLATE FOR INVESTIGATION REPORT"

    table.cell(3, 0).text = "Document No:\nFORM-GMP-QA-0504"
    table.cell(3, 1).text = "Version No.:\n3.0"
    table.cell(3, 2).text = "Effective date:\n04-Jul-2024"

    doc.add_paragraph()


def _add_section_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(11)


def _add_field_content(doc: Document, field_value: str, placeholder: str = "") -> None:
    if field_value:
        doc.add_paragraph(field_value)
    else:
        p = doc.add_paragraph(placeholder or "[TO BE COMPLETED]")
        p.runs[0].italic = True


def _field_val(session: Session, field_id: str) -> str:
    f = session.fields.get(field_id)
    return f.value if f and f.value else ""


def generate_docx(session: Session) -> io.BytesIO:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_header(doc)

    title = doc.add_paragraph("INVESTIGATION REPORT")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    ref = _field_val(session, "document_ref_number") or "<<Document Reference Number>>"
    doc.add_paragraph(f"DOCUMENT REFERENCE NO.: {ref}")
    doc.add_paragraph(f"DATE OF INITIATION: {datetime.now().strftime('%d-%b-%Y')}")
    doc.add_paragraph()

    # Section 1
    _add_section_heading(doc, "1", "SOURCE OF NON-CONFORMITY/UNEXPECTED OUTCOME")
    source_table = doc.add_table(rows=2, cols=3)
    source_table.cell(0, 0).text = "Source of Non-Conformity"
    source_table.cell(0, 1).text = "Category of Non-Conformity"
    source_table.cell(0, 2).text = "Document Reference Number"
    source_table.cell(1, 0).text = _field_val(session, "source_type") or "Deviation"
    source_table.cell(1, 1).text = _field_val(session, "source_category") or "<<Category>>"
    source_table.cell(1, 2).text = _field_val(session, "document_ref_number") or "<<PR#>>"
    doc.add_paragraph()

    # Section 2
    _add_section_heading(doc, "2", "DESCRIPTION OF THE NON-CONFORMITY/UNEXPECTED OUTCOME")
    doc.add_paragraph("2.1. Description:", style="List Number")
    _add_field_content(doc, _field_val(session, "description"))
    doc.add_paragraph("2.2. Reference Documents/Instruments/Equipment/Products/Others:", style="List Number")
    _add_field_content(doc, _field_val(session, "reference_docs"))
    doc.add_paragraph("2.3. Initiator Name and Department:", style="List Number")
    _add_field_content(doc, _field_val(session, "initiator"))

    # Section 3
    _add_section_heading(doc, "3", "PRE-EVALUATION OF THE NON-CONFORMITY/UNEXPECTED OUTCOME")
    doc.add_paragraph("3.1. Initial/Preliminary Impact Assessment:")
    _add_field_content(doc, _field_val(session, "preliminary_impact"))
    doc.add_paragraph("3.2. Immediate Actions or Correction:")
    _add_field_content(doc, _field_val(session, "immediate_actions"))

    # Section 4
    _add_section_heading(doc, "4", "HISTORICAL CHECK")
    hist_table = doc.add_table(rows=5, cols=2)
    hist_table.cell(0, 0).text = "Reference of relevant data source"
    hist_table.cell(0, 1).text = _field_val(session, "historical_data_source") or "[TO BE COMPLETED]"
    hist_table.cell(1, 0).text = "Time frame of review"
    hist_table.cell(1, 1).text = _field_val(session, "historical_timeframe") or "[TO BE COMPLETED]"
    hist_table.cell(2, 0).text = "Justification (if < 12 months)"
    hist_table.cell(2, 1).text = _field_val(session, "historical_justification") or "N/A"
    hist_table.cell(3, 0).text = "Similar in nature (Earlier PR numbers)"
    hist_table.cell(3, 1).text = _field_val(session, "historical_similar_events") or "No similar events identified."
    doc.add_paragraph()

    # Section 5
    _add_section_heading(doc, "5", "INVESTIGATION")
    doc.add_paragraph("5.1 Investigation Team:")
    team_val = _field_val(session, "investigation_team")
    if team_val:
        doc.add_paragraph(team_val)
    else:
        team_table = doc.add_table(rows=2, cols=4)
        for i, h in enumerate(["Name", "Employee ID", "Designation", "Department"]):
            team_table.cell(0, i).text = h
        team_table.cell(1, 0).text = "[TO BE COMPLETED]"

    doc.add_paragraph("5.2 Sequence of Events:")
    _add_field_content(doc, _field_val(session, "sequence_of_events"))

    doc.add_paragraph("5.3 Root Cause Analysis:")
    rca_tools = [
        ("5.3.1", "GEMBA WALK (if applicable)", "gemba_walk"),
        ("5.3.2", "PROCESS MAPPING AND GAP ANALYSIS (if applicable)", "process_mapping"),
        ("5.3.3", "BRAINSTORMING (if applicable)", "brainstorming"),
        ("5.3.4", "Fish bone analysis (if applicable)", "fishbone"),
        ("5.3.5", "Fault Tree Analysis (if applicable)", "fault_tree"),
        ("5.3.6", "Pareto Analysis (if applicable)", "pareto"),
        ("5.3.7", "5-WHY Analysis (if applicable)", "five_why"),
    ]
    for sec_num, sec_title, field_id in rca_tools:
        doc.add_paragraph(f"{sec_num} {sec_title}:")
        val = _field_val(session, field_id)
        if val:
            doc.add_paragraph(val)

    doc.add_paragraph("5.4 Summary of data and documents reviewed:")
    _add_field_content(doc, _field_val(session, "data_summary"))

    hist_rc = _field_val(session, "historical_root_cause_eval")
    if hist_rc:
        doc.add_paragraph("5.5 Evaluation of root cause with historical check findings:")
        doc.add_paragraph(hist_rc)

    doc.add_paragraph("5.6 List of root causes and contributing factors:")
    _add_field_content(doc, _field_val(session, "root_causes"))

    # Section 6
    _add_section_heading(doc, "6", "IMPACT ASSESSMENT")
    impact_fields = [
        ("6.1", "Impact on the quality of the product", "impact_product_quality"),
        ("6.2", "Impact on other products manufactured/analysed in same facility/lab", "impact_other_products"),
        ("6.3", "Impact on other batches, equipment, facility, systems, documents", "impact_other_batches"),
        ("6.4", "Impact on QMS, compliance and regulatory submission documents", "impact_qms_regulatory"),
        ("6.5", "Impact on validated state of product/process/method/stability", "impact_validated_state"),
        ("6.6", "Impact on PM, calibration or qualification status", "impact_pm_calibration"),
    ]
    for sec_num, label, field_id in impact_fields:
        doc.add_paragraph(f"{sec_num} {label}:")
        _add_field_content(doc, _field_val(session, field_id))

    # Section 7
    _add_section_heading(doc, "7", "CORRECTIVE AND PREVENTIVE ACTIONS")
    capa_sections = [
        ("Corrections", "corrections"),
        ("Corrective Actions", "corrective_actions"),
        ("Preventive Actions", "preventive_actions"),
    ]
    for capa_title, field_id in capa_sections:
        doc.add_paragraph(f"{capa_title}:")
        val = _field_val(session, field_id)
        if val:
            cleaned = val
            for noise in ["Sl. No.", "Correction ", "Corrective Action ", "Preventive Action ",
                          "Responsibility (Department)", "Target Timeline", "Target Time line"]:
                cleaned = cleaned.replace(noise, "")
            lines = [l for l in cleaned.split("\n") if l.strip()]
            doc.add_paragraph("\n".join(lines))
        else:
            p = doc.add_paragraph("[TO BE COMPLETED]")
            p.runs[0].italic = True

    # Section 8
    _add_section_heading(doc, "8", "CONCLUSION")
    _add_field_content(doc, _field_val(session, "conclusion"))

    # Section 10
    _add_section_heading(doc, "10", "ABBREVIATIONS")
    abbr = _field_val(session, "abbreviations")
    if abbr:
        doc.add_paragraph(abbr)
    else:
        doc.add_paragraph("[To be populated based on report content]")

    # Section 11
    _add_section_heading(doc, "11", "INVESTIGATION TEAM SIGNATURES")
    sig_table = doc.add_table(rows=5, cols=4)
    headers = ["", "Name", "Designation", "Department"]
    roles = ["Prepared by (User Department)", "Reviewed by (HOD/Designee)",
             "Reviewed by (Investigation Team)", "Approved by (QA)"]
    for i, h in enumerate(headers):
        sig_table.cell(0, i).text = h
    for i, role in enumerate(roles):
        sig_table.cell(i + 1, 0).text = role

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
