"""Extract text from PDF / DOCX / DOC / plaintext uploads.

Ported from Devio knowledge/loader.py. Used both at runtime (user uploads) and
at startup (loading example reports for style matching).
"""

from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path


# Strip C0 control bytes that survive tools like textutil (smart quotes get
# mangled, some docs embed form-feed / vertical-tab). Keep printable + TAB +
# LF + CR. Removes anything else so downstream JSON serialization, LLM prompts,
# and DOCX generation all stay clean.
_CONTROL_BYTE_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _sanitize(text: str) -> str:
    if not text:
        return text
    return _CONTROL_BYTE_RE.sub("", text)


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def extract_doc(path: Path) -> str:
    """Legacy .doc — try antiword, fall back to macOS textutil."""
    for cmd in (["antiword", str(path)],
                ["textutil", "-convert", "txt", "-stdout", str(path)]):
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except FileNotFoundError:
            continue
    return f"[Could not extract text from {path.name} — install antiword or run on macOS with textutil]"


def extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


async def extract_uploaded_file(upload_file) -> tuple[str, str]:
    """FastAPI UploadFile → (filename, extracted text)."""
    content = await upload_file.read()
    filename = upload_file.filename or "unknown"
    suffix = Path(filename).suffix.lower()

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        if suffix == ".pdf":
            text = extract_pdf(tmp_path)
        elif suffix == ".docx":
            text = extract_docx(tmp_path)
        elif suffix == ".doc":
            text = extract_doc(tmp_path)
        elif suffix in (".txt", ".csv", ".log"):
            text = content.decode("utf-8", errors="replace")
        else:
            text = f"[Unsupported file type: {suffix}]"
    finally:
        tmp_path.unlink(missing_ok=True)

    return filename, _sanitize(text)
